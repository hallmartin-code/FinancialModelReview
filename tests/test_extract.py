"""The extraction request shape, checked without touching the network."""

from __future__ import annotations

import json
from types import SimpleNamespace
from typing import Any

import pytest
from reportlab.pdfgen import canvas

from deckscan.extract import claude as claude_module
from deckscan.extract.schema import extraction_schema
from deckscan.extract.source import (
    MAX_TEXT_CHARS,
    _fair_shares,
    prepare_deck,
    prepare_model,
)


def _walk_objects(node: Any) -> list[dict[str, Any]]:
    found: list[dict[str, Any]] = []
    if isinstance(node, dict):
        if node.get("type") == "object":
            found.append(node)
        for value in node.values():
            found.extend(_walk_objects(value))
    elif isinstance(node, list):
        for item in node:
            found.extend(_walk_objects(item))
    return found


def test_schema_is_closed_and_fully_required(settings, fields):
    """Structured outputs reject open objects and partial `required` lists."""
    schema = extraction_schema(settings, fields)
    for obj in _walk_objects(schema):
        assert obj.get("additionalProperties") is False
        assert set(obj.get("required", [])) == set(obj.get("properties", {}))


def test_schema_asks_for_every_template_field(settings, fields):
    schema = extraction_schema(settings, fields)
    assert set(schema["properties"]["narrative"]["properties"]) == set(fields)


def test_schema_metric_names_are_the_configured_ones(settings, fields):
    schema = extraction_schema(settings, fields)
    names = schema["properties"]["metrics"]["items"]["properties"]["name"]["enum"]
    assert set(names) == set(settings.metric_aliases)


# --- source preparation ------------------------------------------------------


@pytest.fixture
def pdf_path(tmp_path):
    path = tmp_path / "deck.pdf"
    pdf = canvas.Canvas(str(path))
    pdf.drawString(72, 720, "Northwind Analytics")
    pdf.showPage()
    pdf.drawString(72, 720, "FY2025 revenue $1.2M")
    pdf.save()
    return path


def test_pdf_is_sent_as_a_native_document_block(pdf_path, settings):
    content = prepare_deck(pdf_path, settings)
    assert content.blocks[0]["type"] == "document"
    assert content.blocks[0]["source"]["media_type"] == "application/pdf"
    assert any("native PDF" in note for note in content.methodology)


def test_word_deck_is_flattened_with_page_markers(tmp_path, settings):
    from docx import Document

    path = tmp_path / "deck.docx"
    document = Document()
    document.add_paragraph("Northwind Analytics")
    document.add_paragraph("FY2025 revenue $1.2M")
    document.save(str(path))

    content = prepare_deck(path, settings)
    assert not content.failed
    assert content.inline_text is not None
    assert "[p.1]" in content.inline_text
    assert "Northwind Analytics" in content.inline_text


def test_csv_model_keeps_cell_references(tmp_path):
    path = tmp_path / "model.csv"
    path.write_text("Period,Revenue\nFY2025,1200000\n", encoding="utf-8")
    content = prepare_model(path)
    assert content.inline_text is not None
    assert "A1=Period" in content.inline_text
    assert "B2=1200000" in content.inline_text


def test_unreadable_source_degrades_instead_of_raising(tmp_path, settings):
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"not a pdf at all")
    content = prepare_deck(path, settings)
    assert content.failed
    assert content.methodology


# --- the API call ------------------------------------------------------------


class _FakeStream:
    def __init__(self, message: Any) -> None:
        self._message = message

    def __enter__(self) -> _FakeStream:
        return self

    def __exit__(self, *exc: object) -> None:
        return None

    def get_final_message(self) -> Any:
        return self._message


def _fake_client(captured: dict[str, Any], message: Any) -> Any:
    def stream(**kwargs: Any) -> _FakeStream:
        captured.update(kwargs)
        return _FakeStream(message)

    return SimpleNamespace(beta=SimpleNamespace(messages=SimpleNamespace(stream=stream)))


def _message(text: str, stop_reason: str = "end_turn") -> Any:
    return SimpleNamespace(
        content=[SimpleNamespace(type="text", text=text)],
        stop_reason=stop_reason,
        stop_details=None,
    )


def test_request_uses_opus_5_structured_output_and_fallbacks(
    monkeypatch, pdf_path, settings, fields
):
    captured: dict[str, Any] = {}
    payload = {"narrative": {}, "series": [], "metrics": [], "claims": {}}
    monkeypatch.setattr(
        claude_module, "get_client", lambda: _fake_client(captured, _message(json.dumps(payload)))
    )

    content = prepare_deck(pdf_path, settings)
    result = claude_module.extract(content, settings, fields)

    assert result == payload
    assert captured["model"] == "claude-opus-5"
    assert captured["betas"] == [claude_module.FALLBACK_BETA]
    assert captured["fallbacks"] == "default"
    assert captured["output_config"]["format"]["type"] == "json_schema"
    assert captured["output_config"]["effort"] == claude_module.EFFORT
    blocks = captured["messages"][0]["content"]
    assert blocks[0]["type"] == "document"
    assert blocks[-1]["type"] == "text"


def test_refusal_becomes_an_extraction_error(monkeypatch, pdf_path, settings, fields):
    monkeypatch.setattr(
        claude_module,
        "get_client",
        lambda: _fake_client({}, _message("", stop_reason="refusal")),
    )
    content = prepare_deck(pdf_path, settings)
    with pytest.raises(claude_module.ExtractionError, match="declined"):
        claude_module.extract(content, settings, fields)


def test_truncation_becomes_an_extraction_error(monkeypatch, pdf_path, settings, fields):
    monkeypatch.setattr(
        claude_module,
        "get_client",
        lambda: _fake_client({}, _message("{", stop_reason="max_tokens")),
    )
    content = prepare_deck(pdf_path, settings)
    with pytest.raises(claude_module.ExtractionError, match="truncated"):
        claude_module.extract(content, settings, fields)


def test_missing_key_is_a_clear_error(settings, fields, pdf_path):
    content = prepare_deck(pdf_path, settings)
    with pytest.raises(claude_module.ExtractionError, match="ANTHROPIC_API_KEY"):
        claude_module.extract(content, settings, fields)


# --- workbooks: every tab, not just the first ---------------------------------


def _workbook(tmp_path, name, tabs, hidden=(), formulas=None):
    """Build a workbook from ``{title: [row, ...]}`` and return its path."""
    from openpyxl import Workbook

    path = tmp_path / name
    book = Workbook()
    for index, (title, rows) in enumerate(tabs.items()):
        sheet = book.active if index == 0 else book.create_sheet()
        sheet.title = title
        for row in rows:
            sheet.append(row)
        if title in hidden:
            sheet.sheet_state = "hidden"
    for title, cells in (formulas or {}).items():
        sheet = book.create_sheet(title)
        for ref, formula in cells.items():
            sheet[ref] = formula
    book.save(str(path))
    return path


def test_every_tab_reaches_the_request(tmp_path):
    """A model's numbers live on later tabs as often as the first one."""
    path = _workbook(
        tmp_path,
        "model.xlsx",
        {
            "Cover": [["Northwind Analytics"]],
            "P&L": [["", "FY2024"], ["Revenue", 1200000]],
            "Cash Flow": [["Ending cash", 800000]],
            "Cap Table": [["Founders", 0.62]],
        },
    )
    content = prepare_model(path)
    text = content.inline_text or ""
    for title in ("Cover", "P&L", "Cash Flow", "Cap Table"):
        assert f"[sheet {title}]" in text
    assert "800000" in text and "0.62" in text


def test_the_methodology_names_every_tab_it_read(tmp_path):
    """'Did it read my tabs?' has to be answerable from the report itself."""
    path = _workbook(
        tmp_path,
        "model.xlsx",
        {"Cover": [["Northwind"]], "P&L": [["Revenue", 1200000]], "Notes": []},
    )
    note = " ".join(prepare_model(path).methodology)
    assert "Cover" in note and "P&L" in note
    assert "Notes" in note  # an empty tab is named as empty, not silently dropped


def test_hidden_tabs_are_read_and_disclosed(tmp_path):
    path = _workbook(
        tmp_path,
        "model.xlsx",
        {"P&L": [["Revenue", 1200000]], "Backup": [["CAC", 4200]]},
        hidden=("Backup",),
    )
    content = prepare_model(path)
    assert "4200" in (content.inline_text or "")
    assert any("Hidden tabs" in note and "Backup" in note for note in content.methodology)


def test_no_tab_is_dropped_when_the_workbook_exceeds_the_budget(tmp_path):
    """The bug this guards: tab 1 eats the budget and later tabs vanish unread."""
    titles = [f"Tab{n:02d}" for n in range(1, 13)]
    rows = [
        [f"Line item {r} with a fairly long label", 1234567, 2345678, 3456789] for r in range(900)
    ]
    path = _workbook(tmp_path, "big.xlsx", dict.fromkeys(titles, rows))

    content = prepare_model(path)
    text = content.inline_text or ""
    assert len(text) > MAX_TEXT_CHARS // 2, "the budget should actually be under pressure"
    assert len(text) <= MAX_TEXT_CHARS, "the hard ceiling still holds"
    missing = [t for t in titles if f"[sheet {t}]" not in text]
    assert not missing, f"tabs dropped entirely: {missing}"
    assert any("clipped" in note for note in content.methodology)


def test_one_huge_tab_does_not_clip_the_small_ones(tmp_path):
    """Cash Flow is two rows; a 9000-row Detail tab must not cost it its place."""
    huge = [[f"Line item {r} with a fairly long label", 1234567, 2345678] for r in range(9000)]
    path = _workbook(
        tmp_path,
        "lopsided.xlsx",
        {
            "Detail": huge,
            "Cash Flow": [["Ending cash", 800000]],
            "Use of Funds": [["Engineering", 1500000]],
        },
    )
    content = prepare_model(path)
    text = content.inline_text or ""
    assert "Ending cash" in text and "1500000" in text
    assert "[sheet Detail]" in text
    clipped = [note for note in content.methodology if "clipped" in note]
    assert clipped and "Detail" in clipped[0]
    assert "Cash Flow" not in clipped[0], "a small tab should never need clipping"


def test_formula_tabs_with_no_cached_values_are_explained(tmp_path):
    """Otherwise a tab of live formulas looks like a tab that states nothing."""
    path = _workbook(
        tmp_path,
        "model.xlsx",
        {"P&L": [["Revenue", 1200000]]},
        formulas={"Calc": {"A1": "Gross margin", "B1": "=1-0.35"}},
    )
    content = prepare_model(path)
    note = " ".join(content.methodology)
    assert "no cached results" in note and "Calc" in note
    assert "re-save" in note


def test_an_all_empty_workbook_fails_with_its_tab_count(tmp_path):
    path = _workbook(tmp_path, "empty.xlsx", {"Sheet1": [], "Sheet2": []})
    content = prepare_model(path)
    assert content.failed
    assert "2 tab(s)" in " ".join(content.methodology)


def test_xlsm_workbooks_read_the_same_way(tmp_path):
    path = _workbook(tmp_path, "model.xlsm", {"P&L": [["Revenue", 1200000]]})
    content = prepare_model(path)
    assert not content.failed
    assert "[sheet P&L]" in (content.inline_text or "")


def test_fair_shares_gives_every_tab_room():
    """Small entries keep their full size; the surplus goes to the large ones."""
    shares = _fair_shares([10, 10, 1000], 300)
    assert shares[0] == 10 and shares[1] == 10
    assert shares[2] == 280
    assert sum(shares) <= 300
